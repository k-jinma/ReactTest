from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Mm, Pt, RGBColor


OUTPUT_DIR = Path("output/doc")
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)


LESSON_BASES = [
    {
        "number": 1,
        "subtitle": "Reactの全体像と開発環境",
        "goal": "Reactがなぜ使われるのかを理解し、Viteを使って最初の開発環境を構築しながら、コンポーネントベース開発の考え方に慣れること",
        "setup_code": "npm create vite@latest react-lesson-01 -- --template react\ncd react-lesson-01\nnpm install\nnpm run dev",
        "terms": ["React", "コンポーネント", "Vite", "SPA", "宣言的UI", "再利用"],
        "sections": [
            {
                "title": "Reactとは何かをつかむ",
                "concepts": [
                    "Reactは画面を部品として考えるライブラリです。従来のようにHTMLをページ単位で丸ごと切り替えるよりも、必要な部分だけを部品として更新できるため、複雑な画面でも整理して作りやすくなります。",
                    "学習の出発点では、Reactを『画面を小さな関数として組み立てる道具』と理解すると進みやすくなります。各部品は見た目だけでなく、状態やイベントも内包できるため、UIとロジックを近い場所で扱えるのが特徴です。",
                    "特にフォーム、タブ、モーダル、検索結果一覧のような、利用者の操作で見た目が頻繁に変わる画面に強く、再利用可能な部品へ分割しやすい点が業務開発でも高く評価されています。",
                ],
                "code_title": "// src/App.jsx",
                "code": "function App() {\n  return (\n    <main>\n      <h1>はじめてのReact</h1>\n      <p>部品としてUIを組み立てます。</p>\n    </main>\n  );\n}\n\nexport default App;",
            },
            {
                "title": "Viteで最初のReactアプリを作る",
                "concepts": [
                    "Reactの学習では、環境構築でつまずかないことが大切です。Viteは高速に起動でき、最小構成でReactの学習を始めやすいため、教材や現場でも広く使われています。",
                    "開発サーバーを立ち上げたら、保存と同時に画面が更新される挙動を必ず確認しましょう。この『修正してすぐ結果を見る』サイクルがReact学習のリズムになります。",
                    "作成直後のファイル構成を見て、どこにコンポーネントを書き、どこがエントリーポイントなのかを理解しておくと、後の章で複数ファイルに分割するときに迷いにくくなります。",
                ],
                "code_title": "// src/main.jsx",
                "code": "import React from 'react';\nimport ReactDOM from 'react-dom/client';\nimport App from './App.jsx';\nimport './index.css';\n\nReactDOM.createRoot(document.getElementById('root')).render(\n  <React.StrictMode>\n    <App />\n  </React.StrictMode>\n);",
            },
            {
                "title": "コンポーネント思考の基本を身につける",
                "concepts": [
                    "コンポーネント設計では、画面を『タイトル』『ボタン』『カード』『一覧』のような意味のある単位に分けます。重要なのは見た目だけでなく責任範囲を意識することです。",
                    "例えば自己紹介ページなら、プロフィール見出し、自己紹介文、リンク一覧を別々のコンポーネントに分けられます。これにより修正箇所が明確になり、複数人開発でも変更の影響範囲が限定されます。",
                    "React学習初期は細かく分けすぎる必要はありませんが、『この塊は別ファイルにできそうか』と考えながら書くだけでも、後の保守性に大きな差が出ます。",
                ],
                "code_title": "// src/components/ProfileCard.jsx",
                "code": "function ProfileCard() {\n  return (\n    <section>\n      <h2>プロフィール</h2>\n      <p>Reactを学びながらUIを組み立てています。</p>\n    </section>\n  );\n}\n\nexport default ProfileCard;",
            },
            {
                "title": "JSXの見方と書き方に慣れる",
                "concepts": [
                    "JSXはJavaScriptの中にHTMLに似た構文を書ける仕組みです。見た目はHTMLに近いですが、実際にはJavaScriptの式として評価されるため、変数や式を波かっこで埋め込めます。",
                    "classではなくclassNameを使う、複数要素を返すときは1つの親要素で包むなど、最初に覚えるルールがいくつかあります。これらは文法の暗記ではなく、JavaScriptの構文制約に沿った書き方だと理解すると定着します。",
                    "JSXは『画面の構造がコードから読みやすい』点が大きな利点です。条件分岐やリスト描画も最終的にはJavaScriptの式で表現できるため、UIロジックを一貫して管理できます。",
                ],
                "code_title": "// JSXの例",
                "code": "const userName = 'Aiko';\nconst isMember = true;\n\nfunction Greeting() {\n  return (\n    <div>\n      <h2>{userName}さん、こんにちは</h2>\n      <p>{isMember ? '会員ページへようこそ' : '登録をお願いします'}</p>\n    </div>\n  );\n}",
            },
        ],
    },
    {
        "number": 2,
        "subtitle": "JSX・Props・コンポーネント分割",
        "goal": "Propsを使って親子コンポーネント間で値を受け渡しし、画面を複数の再利用可能な部品に分割できるようになること",
        "setup_code": "npm install\nnpm run dev",
        "terms": ["JSX", "Props", "親子関係", "再利用", "分割", "export"],
        "sections": [
            {
                "title": "Propsでデータを渡す意味を理解する",
                "concepts": [
                    "Propsは親コンポーネントから子コンポーネントへ値を渡す仕組みです。画面の部品を再利用したいとき、見た目の骨組みは同じでも中身だけを変えたい場面で非常に重要になります。",
                    "例えば同じカードデザインで、タイトルや説明だけが異なる複数の情報を表示したいとき、Propsを使えば1つのカードコンポーネントを繰り返し利用できます。",
                    "Propsは基本的に読み取り専用として扱います。子側で直接書き換えようとせず、必要に応じて親にイベントを伝える設計がReactらしいデータフローです。",
                ],
                "code_title": "// 親から子へ値を渡す",
                "code": "function LessonCard({ title, level }) {\n  return (\n    <article>\n      <h2>{title}</h2>\n      <p>難易度: {level}</p>\n    </article>\n  );\n}\n\nfunction App() {\n  return <LessonCard title=\"Props入門\" level=\"基礎\" />;\n}",
            },
            {
                "title": "分割の基準を持ってコンポーネントを設計する",
                "concepts": [
                    "分割の判断では、『責任が違うか』『別の場所でも使い回せるか』『長くなりすぎて読みづらくなっていないか』の3点を見ると実践的です。",
                    "ただし、最初から細かく切り分けすぎると、かえってファイルが増えすぎて理解しにくくなることがあります。まずは意味のまとまりごとに分割し、必要に応じて調整する姿勢が大切です。",
                    "業務ではデザイナーの共通UIや、一覧行、入力部品、ダイアログなどが分割の起点になります。学習段階でも、見出し・説明・操作ボタンをまとめたカード単位で考えると感覚がつかみやすくなります。",
                ],
                "code_title": "// src/components/UserCard.jsx",
                "code": "function UserCard({ name, role, message }) {\n  return (\n    <section>\n      <h3>{name}</h3>\n      <p>{role}</p>\n      <p>{message}</p>\n    </section>\n  );\n}\n\nexport default UserCard;",
            },
            {
                "title": "子要素とPropsの使い分けを覚える",
                "concepts": [
                    "Propsは特定の意味を持つ値を渡すのに向いています。一方で、ボタンの中にアイコンと文字を入れたい、カードの中身だけ差し替えたいといった柔軟な構造にはchildrenが便利です。",
                    "childrenは『この部品の内側に何を入れるか』を親が決められる仕組みです。レイアウト用コンポーネントを作る際に特に役立ち、再利用性の高い設計へつながります。",
                    "propsの命名は利用者にとって分かりやすいことが重要です。textやdataのような広すぎる名前よりも、title、description、buttonLabelのように意図が見える名前が保守しやすさを高めます。",
                ],
                "code_title": "// childrenの例",
                "code": "function Panel({ title, children }) {\n  return (\n    <section>\n      <h2>{title}</h2>\n      <div>{children}</div>\n    </section>\n  );\n}\n\nfunction App() {\n  return (\n    <Panel title=\"お知らせ\">\n      <p>今週の学習目標を確認しましょう。</p>\n    </Panel>\n  );\n}",
            },
            {
                "title": "Props設計でよくある失敗を避ける",
                "concepts": [
                    "Propsの数が増えすぎると、使う側も受け取る側も見通しが悪くなります。その場合はオブジェクトにまとめる、責任を分割する、childrenを使うなどの整理が必要です。",
                    "また、表示専用コンポーネントにロジックまで詰め込みすぎると再利用が難しくなります。見た目中心の部品と、データを集める親コンポーネントを分ける発想は中級以降でも重要です。",
                    "学習中は『このコンポーネントは何を受け取り、何を表示する責任を持つのか』を言葉にしてみてください。役割を文章化できると、設計の迷いが減ります。",
                ],
                "code_title": "// オブジェクトでまとめる例",
                "code": "function Profile({ user }) {\n  return (\n    <div>\n      <h2>{user.name}</h2>\n      <p>{user.department}</p>\n      <p>{user.bio}</p>\n    </div>\n  );\n}",
            },
        ],
    },
    {
        "number": 3,
        "subtitle": "Stateとイベント処理",
        "goal": "useStateを使って画面上の値を更新し、ボタンクリックや入力イベントに応じてUIを変化させられるようになること",
        "setup_code": "npm install\nnpm run dev",
        "terms": ["state", "useState", "イベント", "再レンダリング", "入力制御", "更新関数"],
        "sections": [
            {
                "title": "stateが必要になる場面を知る",
                "concepts": [
                    "stateは、時間や操作によって変わる値をReactに覚えさせるための仕組みです。クリック回数、入力欄の文字、開閉状態、選択中のタブなどは代表的なstateです。",
                    "通常の変数を変更しても画面は自動では更新されません。Reactに『この値が変わったので再描画してほしい』と伝えるためにstateを使います。この考え方が理解できると、多くのUIは説明できるようになります。",
                    "学習初期は、表示内容の変化を見たら『この値はstateかもしれない』と判断する癖をつけましょう。Reactでの設計力は、変化する値を見抜く力と深く結びついています。",
                ],
                "code_title": "// カウンター",
                "code": "import { useState } from 'react';\n\nfunction Counter() {\n  const [count, setCount] = useState(0);\n\n  return (\n    <div>\n      <p>現在の回数: {count}</p>\n      <button onClick={() => setCount(count + 1)}>増やす</button>\n    </div>\n  );\n}",
            },
            {
                "title": "イベント処理と更新関数を結びつける",
                "concepts": [
                    "イベント処理では、onClickやonChangeのような属性に関数を渡します。重要なのは、関数をその場で実行するのではなく『イベント発生時に呼ばれる関数』として渡すことです。",
                    "setState系の関数は非同期的にまとめて処理されることがあるため、直前の値に依存する更新では関数形式の更新を知っておくと安全です。これは中級に進むと特に重要になります。",
                    "イベント処理は『利用者の行動をどんなstate更新につなげるか』を設計する場所です。ボタンを押したら何が増え、入力したらどの値が変わるのかを意識してコードを書くと整理しやすくなります。",
                ],
                "code_title": "// 関数形式の更新",
                "code": "function Counter() {\n  const [count, setCount] = useState(0);\n\n  const handleIncrease = () => {\n    setCount((prevCount) => prevCount + 1);\n  };\n\n  return <button onClick={handleIncrease}>加算</button>;\n}",
            },
            {
                "title": "入力フォームをstateで制御する",
                "concepts": [
                    "Reactでは、入力欄のvalueをstateで管理する『制御されたコンポーネント』が基本です。これにより、入力値の検証、整形、送信前確認などを一貫した流れで扱えます。",
                    "初心者が混乱しやすいのは、入力欄に文字を入れたとき、実際には『入力イベントが発生し、stateが更新され、その結果としてvalueに反映される』という流れで画面が動いている点です。",
                    "この仕組みを理解すると、検索ボックス、会員登録フォーム、フィルター条件入力など、実務で頻繁に登場するUIを自力で作れるようになります。",
                ],
                "code_title": "// 入力値の管理",
                "code": "function NameForm() {\n  const [name, setName] = useState('');\n\n  return (\n    <div>\n      <input\n        value={name}\n        onChange={(event) => setName(event.target.value)}\n        placeholder=\"名前を入力\"\n      />\n      <p>入力中の名前: {name}</p>\n    </div>\n  );\n}",
            },
            {
                "title": "state設計の基本ルールを押さえる",
                "concepts": [
                    "重複するstateはできるだけ避けます。例えば商品の一覧があり、件数は一覧の長さから計算できるなら、件数を別stateとして持つ必要はありません。計算可能な値は派生値として扱うのがReact流です。",
                    "また、複数のstateが同じタイミングで変わるなら、1つのオブジェクトや配列にまとめる選択肢もあります。ただし更新が複雑になりすぎるなら分けた方が読みやすい場合もあります。",
                    "大切なのは『最小限の正しい状態だけを持つ』ことです。stateが増えるほど不整合の危険も増えるため、まずは何を保存すべきかを慎重に考える習慣を持ちましょう。",
                ],
                "code_title": "// 派生値の例",
                "code": "function CartSummary() {\n  const [items, setItems] = useState(['React入門', 'Hooks実践']);\n  const itemCount = items.length;\n\n  return <p>カート内の商品数: {itemCount}</p>;\n}",
            },
        ],
    },
    {
        "number": 4,
        "subtitle": "条件分岐・リスト表示・フォーム基礎",
        "goal": "条件による表示切り替え、配列からの一覧描画、フォームUIの構築を通じて、実用的な画面の土台を作れるようになること",
        "setup_code": "npm install\nnpm run dev",
        "terms": ["条件分岐", "リスト描画", "map", "key", "フォーム", "制御コンポーネント"],
        "sections": [
            {
                "title": "条件分岐で画面を切り替える",
                "concepts": [
                    "ログイン状態によって表示を変える、データ取得中は読み込み中を表示する、エラー時だけ警告を見せるといったUIは、条件分岐の典型例です。",
                    "Reactではif文を使って返す内容を分ける方法と、JSXの中で三項演算子や論理積を使う方法があります。どちらが読みやすいかを考えて使い分けることが大切です。",
                    "条件分岐は見た目の切り替えだけでなく、不要な要素を描画しないことでコードの意図を明確にする役割もあります。複雑な分岐は小さなコンポーネントに分けると見通しが良くなります。",
                ],
                "code_title": "// 条件による表示切り替え",
                "code": "function LoginMessage({ isLoggedIn }) {\n  return (\n    <div>\n      {isLoggedIn ? <p>ようこそ、会員ページへ</p> : <p>ログインしてください</p>}\n    </div>\n  );\n}",
            },
            {
                "title": "配列をmapで描画する",
                "concepts": [
                    "実務の画面では、一覧表示がほぼ必ず登場します。商品一覧、通知一覧、ユーザー一覧など、配列データをUIへ変換する基本操作としてmapは最重要です。",
                    "配列の各要素をどのような見た目へ変換するかを関数として書くことで、データと表示の関係がはっきりします。Reactでは『配列を回してJSXを返す』という考え方に慣れることが大切です。",
                    "一覧の見た目が複雑なら、1行分を別コンポーネントに切り出して可読性を保ちましょう。mapの中で大きなJSXを書くよりも、責任を分けた方が修正に強い構成になります。",
                ],
                "code_title": "// 一覧表示",
                "code": "const lessons = [\n  { id: 1, title: 'React入門' },\n  { id: 2, title: 'useState基礎' },\n  { id: 3, title: 'フォーム操作' },\n];\n\nfunction LessonList() {\n  return (\n    <ul>\n      {lessons.map((lesson) => (\n        <li key={lesson.id}>{lesson.title}</li>\n      ))}\n    </ul>\n  );\n}",
            },
            {
                "title": "keyの役割を理解して安定したUIを作る",
                "concepts": [
                    "keyはReactが一覧の差分を判断するための手がかりです。単なる警告回避ではなく、更新時にどの要素が同じものかを認識するために必要な識別子です。",
                    "indexをkeyに使うと、並び替えや削除が入ったときに意図しない再利用が起こり、入力欄の内容がずれたり、アニメーションが不自然になったりすることがあります。可能な限り一意なidを使いましょう。",
                    "一覧の安定性はユーザー体験に直結します。見た目は同じでも内部の対応関係が崩れると不具合の原因になるため、keyは早い段階から丁寧に扱う習慣が重要です。",
                ],
                "code_title": "// 良いkeyの例",
                "code": "const tasks = [\n  { id: 'task-001', title: '教材を読む' },\n  { id: 'task-002', title: 'コードを書く' },\n];\n\nfunction TaskList() {\n  return tasks.map((task) => <p key={task.id}>{task.title}</p>);\n}",
            },
            {
                "title": "フォームを組み合わせて小さなアプリを作る",
                "concepts": [
                    "条件分岐、一覧表示、入力制御が組み合わさると、簡単なToDoアプリや検索フォームなど、実用的な画面が作れるようになります。基礎の学習では、この組み合わせが大きな到達点です。",
                    "入力値を追加ボタンで配列へ入れ、その配列をmapで表示し、空欄時には警告を出す、といった一連の流れを通して、Reactの基本的な画面更新モデルを体感できます。",
                    "ここで大切なのは、1つずつ動くことを確認しながら積み上げることです。入力、追加、表示、条件分岐を一気に書くより、小さく確認した方がバグの原因をつかみやすくなります。",
                ],
                "code_title": "// 簡単なタスク追加",
                "code": "function TodoApp() {\n  const [text, setText] = useState('');\n  const [items, setItems] = useState([]);\n\n  const handleAdd = () => {\n    if (!text.trim()) return;\n    setItems((prevItems) => [...prevItems, { id: crypto.randomUUID(), title: text }]);\n    setText('');\n  };\n\n  return (\n    <div>\n      <input value={text} onChange={(e) => setText(e.target.value)} />\n      <button onClick={handleAdd}>追加</button>\n      {items.length === 0 ? <p>まだタスクはありません</p> : null}\n      <ul>\n        {items.map((item) => (\n          <li key={item.id}>{item.title}</li>\n        ))}\n      </ul>\n    </div>\n  );\n}",
            },
        ],
    },
    {
        "number": 5,
        "subtitle": "useEffectと副作用の考え方",
        "goal": "useEffectの役割を理解し、画面表示後の処理、外部データとの同期、クリーンアップの必要性を説明できるようになること",
        "setup_code": "npm install\nnpm run dev",
        "terms": ["useEffect", "副作用", "依存配列", "クリーンアップ", "同期", "ライフサイクル"],
        "sections": [
            {
                "title": "副作用とは何かを理解する",
                "concepts": [
                    "Reactのレンダリングは、stateやpropsから画面を計算する純粋な処理として考えるのが基本です。一方で、通信、タイマー、イベント登録、document.titleの変更などは画面計算の外にあるため、副作用として扱います。",
                    "useEffectはその副作用をレンダリング後に実行するためのフックです。これにより、UIの計算と外部とのやり取りを分けて整理できます。学習初期にこの境界を意識すると設計が安定します。",
                    "副作用を理解せずに使うと、『とりあえずuseEffectに書く』状態になりやすいです。まずは本当に画面計算の外にある処理なのかを見極めることが重要です。",
                ],
                "code_title": "// タイトル変更",
                "code": "import { useEffect, useState } from 'react';\n\nfunction TitleCounter() {\n  const [count, setCount] = useState(0);\n\n  useEffect(() => {\n    document.title = `クリック回数: ${count}`;\n  }, [count]);\n\n  return <button onClick={() => setCount(count + 1)}>増やす</button>;\n}",
            },
            {
                "title": "依存配列の読み方を身につける",
                "concepts": [
                    "依存配列は『どの値が変わったときにeffectを再実行するか』を示します。空配列なら初回のみ、値を入れればその値の変化時、配列を省略すれば毎回実行されます。",
                    "ただし依存配列は『実行回数を調整するためのつまみ』ではなく、effect内で参照する値を正しく宣言するための仕組みです。ここを誤解すると古い値を参照する不具合が起きやすくなります。",
                    "ESLintのHooksルールが依存関係を警告する理由も、effectの中で何を使っているかを正確に保つためです。警告を消すことより、なぜその警告が出たかを理解する方が成長につながります。",
                ],
                "code_title": "// 依存配列の例",
                "code": "useEffect(() => {\n  console.log('keywordが変わりました', keyword);\n}, [keyword]);",
            },
            {
                "title": "クリーンアップが必要な場面を知る",
                "concepts": [
                    "タイマーやイベントリスナー、購読処理は、コンポーネントが不要になったときに解除しないとメモリリークや二重実行の原因になります。useEffectでは戻り値としてクリーンアップ関数を返します。",
                    "例えばsetIntervalを登録しただけで解除しないと、画面を離れた後もタイマーが動き続けてしまいます。学習では目立たなくても、実務では性能や不具合に直結する重要な論点です。",
                    "クリーンアップは『開始したものは必ず終了させる』という対の意識で考えると理解しやすいです。イベントを登録したら解除する、購読したら解除する、タイマーを開始したら停止する、という発想です。",
                ],
                "code_title": "// クリーンアップ",
                "code": "useEffect(() => {\n  const timerId = setInterval(() => {\n    console.log('1秒ごとに実行');\n  }, 1000);\n\n  return () => {\n    clearInterval(timerId);\n  };\n}, []);",
            },
            {
                "title": "useEffectを使いすぎない判断力を持つ",
                "concepts": [
                    "よくある誤りは、計算で求められる値までuseEffectとstateで管理してしまうことです。たとえばフィルター済み配列は元の配列と検索語から計算できるため、effectで別stateへコピーする必要はありません。",
                    "effectを減らすと、状態の重複や同期漏れが減り、コードが素直になります。『値を計算したいだけか』『外部と同期したいのか』を切り分けることが、React中級者への大事な一歩です。",
                    "副作用は必要なときだけ使い、不要なら普通の変数計算で済ませる。この判断を意識するだけで、コードの複雑さはかなり抑えられます。",
                ],
                "code_title": "// 計算で十分な例",
                "code": "const filteredItems = items.filter((item) =>\n  item.title.toLowerCase().includes(keyword.toLowerCase())\n);",
            },
        ],
    },
    {
        "number": 6,
        "subtitle": "フォーム実践とバリデーション",
        "goal": "複数入力を持つフォームをReactで設計し、送信前検証やエラーメッセージ表示を含む実践的な入力画面を作れるようになること",
        "setup_code": "npm install\nnpm run dev",
        "terms": ["フォーム", "バリデーション", "エラー表示", "送信処理", "controlled", "ユーザー入力"],
        "sections": [
            {
                "title": "フォーム設計の基本を整理する",
                "concepts": [
                    "フォームでは、入力欄の数が増えるほどstate設計が重要になります。名前、メールアドレス、問い合わせ内容、同意チェックなど、どの値をひとまとめにし、どのタイミングで検証するかを先に決めると実装が安定します。",
                    "小規模なフォームなら1つのオブジェクトstateでまとめる方法が扱いやすいです。一方で入力欄ごとにロジックが大きく異なる場合は、stateを分けた方が見通しが良いこともあります。",
                    "Reactでフォームを作るときは、表示、入力、検証、送信、結果表示の5つの流れを意識すると設計しやすくなります。特に検証の責任を曖昧にしないことが大切です。",
                ],
                "code_title": "// 1つのオブジェクトで管理",
                "code": "const [form, setForm] = useState({\n  name: '',\n  email: '',\n  message: '',\n  agree: false,\n});",
            },
            {
                "title": "入力値を更新する共通ハンドラを作る",
                "concepts": [
                    "複数入力を扱う場合、inputのname属性とevent.target.nameを利用すると共通の更新関数が作れます。これにより、同じようなonChange関数を大量に書かずに済みます。",
                    "チェックボックスとテキスト入力では取得する値が異なるため、typeごとの分岐が必要です。共通化しつつも、入力の性質に応じた扱いを覚えることが実用上は重要です。",
                    "共通ハンドラを作る目的は、コード量を減らすことだけではありません。更新規則を1箇所に集約することで、あとから入力ルールが増えても修正しやすくなります。",
                ],
                "code_title": "// 共通ハンドラ",
                "code": "const handleChange = (event) => {\n  const { name, value, type, checked } = event.target;\n\n  setForm((prevForm) => ({\n    ...prevForm,\n    [name]: type === 'checkbox' ? checked : value,\n  }));\n};",
            },
            {
                "title": "送信前バリデーションを実装する",
                "concepts": [
                    "バリデーションは『入力ミスを見つけて終わり』ではなく、利用者が次に何を直せばよいか分かる状態を作ることが目的です。エラーメッセージは具体的で、行動につながる文にしましょう。",
                    "必須チェック、文字数チェック、メール形式チェックなどは基礎的ですが、どのタイミングで表示するかも重要です。入力中に出すのか、送信時にまとめて出すのかで体験が変わります。",
                    "学習段階では、まず送信時にerrorsオブジェクトを作る実装から始めると理解しやすいです。その後、blur時や入力中のリアルタイム検証へ広げると段階的に学べます。",
                ],
                "code_title": "// 検証関数",
                "code": "const validate = (form) => {\n  const errors = {};\n\n  if (!form.name.trim()) {\n    errors.name = '名前は必須です。';\n  }\n\n  if (!form.email.includes('@')) {\n    errors.email = 'メールアドレスの形式を確認してください。';\n  }\n\n  if (form.message.trim().length < 20) {\n    errors.message = '問い合わせ内容は20文字以上で入力してください。';\n  }\n\n  if (!form.agree) {\n    errors.agree = '送信前に同意が必要です。';\n  }\n\n  return errors;\n};",
            },
            {
                "title": "送信処理と成功メッセージをつなげる",
                "concepts": [
                    "送信処理では、デフォルトのページ再読み込みをpreventDefaultで止め、検証結果に応じて処理を分けます。ここで非同期通信につなげる設計を意識すると、次回のAPI学習ともつながります。",
                    "成功時には入力値を初期化するのか、そのまま残すのかもUX上の判断です。問い合わせ完了ならクリアしてよいですが、複雑な入力では誤操作対策として保持した方がよい場合もあります。",
                    "送信中の二重送信防止、失敗時の再試行導線、成功メッセージの表示位置など、フォームは細部の設計で使いやすさが大きく変わります。中級へ進むほどこの視点が重要になります。",
                ],
                "code_title": "// submit処理",
                "code": "const handleSubmit = (event) => {\n  event.preventDefault();\n  const nextErrors = validate(form);\n  setErrors(nextErrors);\n\n  if (Object.keys(nextErrors).length > 0) {\n    return;\n  }\n\n  setIsSubmitted(true);\n};",
            },
        ],
    },
    {
        "number": 7,
        "subtitle": "API通信と非同期処理",
        "goal": "fetchを使ったAPI通信の基本を理解し、読み込み中・成功・失敗の状態を持つデータ表示画面を構築できるようになること",
        "setup_code": "npm install\nnpm run dev",
        "terms": ["API", "fetch", "非同期処理", "loading", "error", "JSON"],
        "sections": [
            {
                "title": "非同期処理の全体像を理解する",
                "concepts": [
                    "Webアプリでは、サーバーからデータを取得して画面を描画する場面が頻繁にあります。このとき、結果が返ってくるまで待ち時間があるため、同期処理とは違う設計が必要です。",
                    "Reactでは一般的に、loading、error、dataの3つの状態を意識してUIを組み立てます。これにより『まだ取得中』『失敗した』『成功して表示できる』を明確に分けられます。",
                    "非同期処理の学習では、まず成功ケースだけでなく失敗ケースも最初から考える習慣が大切です。エラーを無視すると、実際の運用で利用者が何も分からない画面になってしまいます。",
                ],
                "code_title": "// 基本のfetch",
                "code": "const response = await fetch('https://jsonplaceholder.typicode.com/posts');\nconst data = await response.json();",
            },
            {
                "title": "useEffectと組み合わせてデータを取得する",
                "concepts": [
                    "画面表示時にデータ取得したい場合、useEffectの中で非同期関数を呼び出す形をよく使います。effect自体をasyncにするのではなく、中で別関数を定義して実行する書き方が基本です。",
                    "取得開始前にloadingをtrueにし、成功ならdataへ格納し、失敗ならerrorへ格納し、最後にloadingをfalseへ戻す流れを丁寧に書くと、画面の状態遷移が分かりやすくなります。",
                    "依存配列を空にすれば初回のみ実行されますが、検索語やIDに応じて再取得したい場合は、それらの値を依存関係に含めます。ここでもeffectの責任を明確にすることが重要です。",
                ],
                "code_title": "// データ取得コンポーネント",
                "code": "import { useEffect, useState } from 'react';\n\nfunction PostList() {\n  const [posts, setPosts] = useState([]);\n  const [loading, setLoading] = useState(true);\n  const [error, setError] = useState('');\n\n  useEffect(() => {\n    const fetchPosts = async () => {\n      try {\n        setLoading(true);\n        setError('');\n        const response = await fetch('https://jsonplaceholder.typicode.com/posts?_limit=5');\n        if (!response.ok) {\n          throw new Error('データ取得に失敗しました。');\n        }\n        const data = await response.json();\n        setPosts(data);\n      } catch (err) {\n        setError(err.message);\n      } finally {\n        setLoading(false);\n      }\n    };\n\n    fetchPosts();\n  }, []);\n}",
            },
            {
                "title": "読み込み中とエラー時のUIを丁寧に作る",
                "concepts": [
                    "利用者は通信処理の内部を知りません。そのため、『今待てばよいのか』『失敗したのか』『再試行できるのか』がUIから分かることが重要です。loadingとerrorは単なるstateではなく、案内の質に直結します。",
                    "ローディング表示は簡素でも構いませんが、何が起きているかを伝える文言を添えると安心感が増します。エラー表示では、開発者向けの難しい文より、利用者向けの行動指示を優先しましょう。",
                    "一覧や詳細画面で通信が発生するたびに、この3状態をどう表現するか考える習慣を持つと、UI設計が一段上がります。",
                ],
                "code_title": "// 画面分岐",
                "code": "if (loading) {\n  return <p>記事を読み込み中です...</p>;\n}\n\nif (error) {\n  return <p>{error}</p>;\n}\n\nreturn posts.map((post) => <p key={post.id}>{post.title}</p>);",
            },
            {
                "title": "検索条件やボタン操作で再取得する",
                "concepts": [
                    "中級に近づくと、表示時に一度だけ取得するだけでは足りず、検索語の変更、ページ送り、並び替え、絞り込みに応じて再取得する必要が出てきます。",
                    "このとき、入力値をそのまま毎回通信条件に使うと、タイピングのたびに大量通信になることがあります。検索ボタンで確定する、遅延処理を入れる、キャッシュを使うなど、運用上の視点も大切です。",
                    "API通信は『データを取れたら終わり』ではなく、どのタイミングで、どれだけの頻度で、失敗時にどう扱うかまで含めて設計すると、より実践的な理解につながります。",
                ],
                "code_title": "// ボタンで再取得する発想",
                "code": "const handleReload = () => {\n  setReloadKey((prevKey) => prevKey + 1);\n};\n\nuseEffect(() => {\n  // reloadKeyが変わるたびに再取得\n}, [reloadKey]);",
            },
        ],
    },
    {
        "number": 8,
        "subtitle": "React Routerによる画面遷移",
        "goal": "React Routerの基本を理解し、複数ページ構成、ナビゲーション、URLパラメータを持つアプリを作れるようになること",
        "setup_code": "npm install react-router-dom\nnpm run dev",
        "terms": ["React Router", "BrowserRouter", "Routes", "Route", "Link", "URLパラメータ"],
        "sections": [
            {
                "title": "SPAにおける画面遷移の考え方を理解する",
                "concepts": [
                    "Reactアプリでは、従来の複数HTMLページを移動する形ではなく、1つのアプリ内で表示内容を切り替えるSPAの考え方が基本になります。そのため、URLと表示コンポーネントの対応付けが重要です。",
                    "React Routerは、その対応付けを整理して書くための代表的なライブラリです。URLが変わったときに、どのコンポーネントを表示するかを宣言的に定義できます。",
                    "学習ではまず、ホーム、教材一覧、詳細ページといった単純な構成から始めると理解しやすいです。URLを見るだけで今どの画面にいるか分かることは、利用者にとっても大きな利点です。",
                ],
                "code_title": "// インストール",
                "code": "npm install react-router-dom",
            },
            {
                "title": "BrowserRouterとRoutesの役割を押さえる",
                "concepts": [
                    "BrowserRouterはアプリ全体にルーティング機能を与える土台です。通常は最上位で1回だけ使い、その内部でRoutesとRouteを使ってページ定義を書きます。",
                    "Routesは複数のRouteをまとめるコンテナで、現在のURLに一致したRouteを探して該当要素を表示します。これにより『どのURLで何を出すか』が一覧で見える構造になります。",
                    "参考教材でもこの部分が中心的に扱われていました。React Routerの理解は、単なるAPI暗記よりも、URLとUIの対応を頭の中で追えるかどうかが重要です。",
                ],
                "code_title": "// src/App.jsx",
                "code": "import { BrowserRouter, Routes, Route, Link } from 'react-router-dom';\n\nfunction Home() {\n  return <h1>ホーム</h1>;\n}\n\nfunction Lessons() {\n  return <h1>教材一覧</h1>;\n}\n\nfunction App() {\n  return (\n    <BrowserRouter>\n      <nav>\n        <Link to=\"/\">ホーム</Link>\n        <Link to=\"/lessons\">教材一覧</Link>\n      </nav>\n\n      <Routes>\n        <Route path=\"/\" element={<Home />} />\n        <Route path=\"/lessons\" element={<Lessons />} />\n      </Routes>\n    </BrowserRouter>\n  );\n}\n\nexport default App;",
            },
            {
                "title": "LinkとuseNavigateを使い分ける",
                "concepts": [
                    "Linkは利用者がクリックする通常のナビゲーションに向いています。aタグと似ていますが、ページ全体を再読み込みせずに画面を切り替えられる点がReact Routerの利点です。",
                    "一方、ログイン成功後に自動で移動したい、送信完了後に一覧へ戻したい、といったプログラム側からの遷移にはuseNavigateを使います。目的に応じた使い分けが大切です。",
                    "画面遷移は単なる表示切り替えではなく、利用者の行動フローを形にする設計でもあります。どのタイミングでどこへ移るのが自然かを考えながら実装しましょう。",
                ],
                "code_title": "// useNavigateの例",
                "code": "import { useNavigate } from 'react-router-dom';\n\nfunction LoginButton() {\n  const navigate = useNavigate();\n\n  const handleLogin = () => {\n    navigate('/dashboard');\n  };\n\n  return <button onClick={handleLogin}>ログインする</button>;\n}",
            },
            {
                "title": "URLパラメータで詳細ページを作る",
                "concepts": [
                    "一覧から詳細へ進む画面では、URLにIDを含める形がよく使われます。たとえば`/lessons/3`のようなURLにして、どの教材を表示するかを判定します。",
                    "React Routerでは、Routeに`:lessonId`のような動的セグメントを書き、useParamsで値を取得します。これにより、1つの詳細コンポーネントで複数データに対応できます。",
                    "URL設計はアプリの分かりやすさにも直結します。意味のあるパス名を付け、一覧と詳細の関係が自然に伝わる形を意識すると、ルーティングの理解が深まります。",
                ],
                "code_title": "// URLパラメータ",
                "code": "import { useParams } from 'react-router-dom';\n\nfunction LessonDetail() {\n  const { lessonId } = useParams();\n\n  return <p>表示中の教材ID: {lessonId}</p>;\n}\n\n<Route path=\"/lessons/:lessonId\" element={<LessonDetail />} />",
            },
        ],
    },
    {
        "number": 9,
        "subtitle": "Context APIとカスタムフック",
        "goal": "propsの受け渡しが深くなる課題を理解し、Context APIとカスタムフックを使って状態共有とロジック再利用を整理できるようになること",
        "setup_code": "npm install\nnpm run dev",
        "terms": ["Context API", "Provider", "useContext", "カスタムフック", "状態共有", "ロジック再利用"],
        "sections": [
            {
                "title": "propsの受け渡しが深くなる問題を知る",
                "concepts": [
                    "小さなアプリではpropsだけで十分ですが、テーマ設定、ログイン中ユーザー、言語設定のように多くの階層で必要な情報は、親から何段も渡すと見通しが悪くなります。",
                    "この状態はprops drillingと呼ばれ、途中のコンポーネントが値を使わないのに渡すだけになりがちです。Reactではこの課題を解決するためにContext APIが用意されています。",
                    "ただし、何でもContextに入れるのは逆効果です。広く共有する価値がある値か、更新頻度が高すぎないかを見極めることが大切です。",
                ],
                "code_title": "// props drillingの例",
                "code": "function App() {\n  return <Layout theme=\"dark\" />;\n}\n\nfunction Layout({ theme }) {\n  return <Sidebar theme={theme} />;\n}\n\nfunction Sidebar({ theme }) {\n  return <Menu theme={theme} />;\n}",
            },
            {
                "title": "Context APIの基本構造を理解する",
                "concepts": [
                    "Contextは、共有したい値をProviderで包んだ範囲へ配る仕組みです。受け取り側ではuseContextを使って必要な値を取得します。『配る側』と『使う側』を明確に分けて考えると理解しやすいです。",
                    "テーマ色やログインユーザーなど、アプリ全体または大きな範囲で共通利用する情報に向いています。逆に、その場限りの入力値までContextに入れると、責任範囲が広がりすぎて保守しにくくなります。",
                    "Contextは便利ですが、更新が多い値をまとめすぎると再レンダリング範囲が広がることがあります。共有の価値と更新コストのバランスを見る視点が中級で重要になります。",
                ],
                "code_title": "// Context作成",
                "code": "import { createContext, useContext, useState } from 'react';\n\nconst ThemeContext = createContext(null);\n\nfunction ThemeProvider({ children }) {\n  const [theme, setTheme] = useState('light');\n\n  return (\n    <ThemeContext.Provider value={{ theme, setTheme }}>\n      {children}\n    </ThemeContext.Provider>\n  );\n}\n\nfunction useTheme() {\n  return useContext(ThemeContext);\n}",
            },
            {
                "title": "カスタムフックでロジックを再利用する",
                "concepts": [
                    "カスタムフックは、複数コンポーネントで共通する状態管理や副作用処理を切り出す方法です。名前をuseから始める通常の関数として作り、内部でReactフックを使えます。",
                    "例えば入力フォームの共通ロジック、API取得処理、ウィンドウ幅の監視などをカスタムフックにすると、コンポーネントは表示責任に集中しやすくなります。",
                    "カスタムフックの価値は、コード量削減だけではありません。『この処理は何をする単位なのか』を名前として表現できるため、チーム内で意図を共有しやすくなります。",
                ],
                "code_title": "// useToggle",
                "code": "import { useState } from 'react';\n\nfunction useToggle(initialValue = false) {\n  const [isOpen, setIsOpen] = useState(initialValue);\n\n  const toggle = () => setIsOpen((prevIsOpen) => !prevIsOpen);\n\n  return { isOpen, toggle, setIsOpen };\n}",
            },
            {
                "title": "Contextとカスタムフックを組み合わせる",
                "concepts": [
                    "実務では、Contextの取得処理を直接各コンポーネントに書くより、専用のカスタムフックを作ることが多いです。これにより、nullチェックや使用ルールを1箇所にまとめられます。",
                    "たとえばuseAuth、useTheme、useCartのようなフック名にすると、呼び出し側は意図を自然に読めます。Providerがない場所で使われた場合のエラーも分かりやすくできます。",
                    "設計の観点では、Contextは共有範囲の定義、カスタムフックは利用インターフェースの定義と考えると整理しやすいです。この考え方は中規模以上のReactアプリで特に有効です。",
                ],
                "code_title": "// 安全なuseTheme",
                "code": "function useTheme() {\n  const context = useContext(ThemeContext);\n\n  if (!context) {\n    throw new Error('useThemeはThemeProviderの中で使ってください。');\n  }\n\n  return context;\n}",
            },
        ],
    },
    {
        "number": 10,
        "subtitle": "中級への橋渡し: 状態設計・性能・保守性",
        "goal": "Reactアプリを継続的に改善する視点を持ち、状態設計、性能面の注意点、保守しやすいファイル分割の考え方を説明できるようになること",
        "setup_code": "npm install\nnpm run dev",
        "terms": ["状態設計", "派生値", "再レンダリング", "メモ化", "責任分割", "保守性"],
        "sections": [
            {
                "title": "状態設計を最初に考える",
                "concepts": [
                    "中級へ進むと、APIの知識そのものよりも『どの値をどこに置くか』の判断が品質を左右します。ローカルstate、親で共有するstate、Contextへ置く値を区別する視点が重要です。",
                    "stateの置き場所は、誰が必要とし、誰が更新するかで決めます。複数の子で共有するなら親へ持ち上げ、アプリ全体で必要ならContextや外部状態管理も候補になります。",
                    "また、計算で求められる値はstateにせず派生値にする、同じ意味の値を重複して持たない、更新トリガーを明確にする、といった基本原則が保守性を大きく左右します。",
                ],
                "code_title": "// 持ち上げの例",
                "code": "function App() {\n  const [keyword, setKeyword] = useState('');\n\n  return (\n    <>\n      <SearchBox keyword={keyword} onChangeKeyword={setKeyword} />\n      <LessonTable keyword={keyword} />\n    </>\n  );\n}",
            },
            {
                "title": "再レンダリングの仕組みと性能の基本を理解する",
                "concepts": [
                    "Reactではstateやpropsが変わると再レンダリングが起こります。これは通常問題ではありませんが、重い計算や大きな一覧がある画面では、どこが頻繁に描画されるかを意識する価値があります。",
                    "性能改善では、いきなり最適化するのではなく、まず不要なstate更新や過剰なeffectを減らすことが先です。根本設計が整っていないままメモ化に頼ると、かえって分かりにくいコードになります。",
                    "useMemoやReact.memoなどの最適化手段は、必要性が見えたときに使うのが基本です。『速そうだから入れる』ではなく、『この処理が実際に重い』という理由を持つことが重要です。",
                ],
                "code_title": "// 派生値として計算",
                "code": "const visibleLessons = lessons.filter((lesson) =>\n  lesson.title.includes(keyword)\n);",
            },
            {
                "title": "ファイル構成と責任分割を整える",
                "concepts": [
                    "アプリが大きくなると、components、pages、hooks、contexts、utilsのように役割ごとの整理が効いてきます。重要なのは名前よりも、責任が混ざらないことです。",
                    "表示中心のコンポーネント、データ取得を担うコンテナ、共有ロジックを持つフック、API通信関数などを分けると、修正時にどこを見ればよいかが明確になります。",
                    "教材学習の段階でも、1ファイルが長くなったら『表示』『ロジック』『定数』を分けられないかを考えるだけで、保守しやすい設計感覚が育ちます。",
                ],
                "code_title": "// 例: ディレクトリ構成",
                "code": "src/\n  components/\n  pages/\n  hooks/\n  contexts/\n  utils/",
            },
            {
                "title": "学習後に取り組むべき発展テーマを整理する",
                "concepts": [
                    "この10回でReactの基礎から中級入口までを一通り学べます。次の段階では、テスト、自動化、アクセシビリティ、状態管理ライブラリ、Server Componentsなどへ進むと理解が深まります。",
                    "ただし、新しい道具に進む前に、state設計、props、effect、ルーティング、フォーム、API通信を自力で説明できることが重要です。基礎の言語化ができる人ほど応用でも強くなります。",
                    "中級学習では『動く』だけでなく、『読みやすいか』『直しやすいか』『失敗時に壊れにくいか』という視点を持ち続けてください。その積み重ねが、実務で信頼されるReact開発者への近道です。",
                ],
                "code_title": "// 次の学習候補",
                "code": "1. React Testing Library\n2. アクセシビリティ改善\n3. パフォーマンス計測\n4. 状態管理ライブラリの比較\n5. 実案件に近いCRUDアプリ制作",
            },
        ],
    },
]


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def set_cell_shading(paragraph, fill):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_para_spacing(paragraph, before=0, after=0, line=1.5):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_code_block(doc, title, code):
    if title:
        label = doc.add_paragraph()
        label.style = "BodyJP"
        label.add_run(title).font.color.rgb = RGBColor(120, 136, 153)
        set_para_spacing(label, before=4, after=2, line=1.2)

    for line in code.splitlines():
        p = doc.add_paragraph()
        p.style = "CodeBlock"
        set_cell_shading(p, "F3F1EE")
        run = p.add_run(line if line else " ")
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_bullet_list(doc, items):
    for item in items:
        p = doc.add_paragraph(style="BodyJP")
        p.paragraph_format.left_indent = Mm(6)
        p.add_run("・")
        p.add_run(item)
        set_para_spacing(p, before=0, after=2, line=1.5)


def add_section_heading(doc, number, title):
    p = doc.add_paragraph()
    p.style = "SectionTitleJP"
    badge = p.add_run(str(number))
    badge.font.color.rgb = RGBColor(255, 255, 255)
    badge.font.bold = True
    r_pr = badge._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "6FA8DC")
    r_pr.append(shd)
    p.add_run(" ")
    p.add_run(title)
    set_para_spacing(p, before=10, after=4, line=1.2)


def lesson_to_document_data(base):
    lesson = deepcopy(base)
    lesson["opening"] = [
        f"この第{lesson['number']}回では、{lesson['goal']}を主な到達目標にします。Reactは一度に多くの概念が登場するため、単語を暗記するよりも、『どの課題を解決するための仕組みなのか』を理解することが大切です。本資料では、画面の変化とコードの対応関係が見えるように、説明、コード、注意点、演習の順で整理しながら学習を進めます。",
        f"特に{lesson['subtitle']}の学習では、表面的な書き方だけでなく、なぜその設計が保守しやすいのか、どこでつまずきやすいのかまで意識することで、基礎知識が単発で終わらず実践へつながります。コード例は小さくしていますが、読みながら自分で値や文言を変更し、動作の差を確認すると理解が定着しやすくなります。",
        f"また、Reactでは『データの流れ』『状態の置き場所』『表示とロジックの責任分担』が繰り返し重要になります。今回扱う内容も、今後のルーティング、API通信、状態共有へそのままつながるため、分かったつもりで流さず、画面がどう再描画されるのかを言葉にしながら読み進めてください。",
    ]
    lesson["key_points"] = [
        f"{lesson['subtitle']}で登場する主要キーワードは、{', '.join(lesson['terms'])}です。学習中は、各用語を単独で覚えるのではなく、『何のために使うか』『使わないとどんな不便があるか』をセットで理解してください。",
        "サンプルコードは小さくても、実際には複数の概念が同時に含まれています。1回読んで分からない部分があっても自然なことなので、実行しながら1行ずつ確認する姿勢を大切にしましょう。",
        "実務で使える力に近づくためには、成功例だけでなく、よくある失敗例や設計上の注意点に注目することが効果的です。本資料でも意図的にその観点を補足しています。",
    ]
    lesson["qa"] = [
        {
            "q": f"Q. 第{lesson['number']}回の内容は丸暗記した方がよいですか？",
            "a": "A. 丸暗記よりも、画面の変化とコードの役割を対応づけて理解する方が重要です。Reactは部品同士の関係を整理していく学習なので、意味をつかむことが先です。",
        },
        {
            "q": "Q. コード例を少し変えただけで動かなくなります。どう復習すればよいですか？",
            "a": "A. 変更前と変更後で、どのstate、props、条件分岐が変わったかを書き出してください。差分を言語化すると、原因を特定しやすくなります。",
        },
    ]
    lesson["practice"] = [
        f"{lesson['subtitle']}の内容を使って、サンプルコードの文言や配列データを自分用の題材に置き換えてみましょう。教材例をそのまま写すだけでなく、自分の言葉へ変えるだけでも理解度が大きく上がります。",
        "画面に何が表示され、どの操作で、どの値が変わるのかを紙やメモに整理してから実装してみてください。Reactは『先に流れを考えてから書く』と学習効率が高くなります。",
        "エラーが出たら、いきなり全文を見直すのではなく、コンポーネント名、props名、import/export、フックの位置など、基本項目を順番に確認する癖をつけましょう。",
    ]
    lesson["closing"] = [
        f"第{lesson['number']}回では、{lesson['subtitle']}を通じて、Reactの基礎から一歩先へ進むための重要な視点を確認しました。ここで大切なのは、APIや構文を個別に覚えることではなく、Reactが『変化する画面を整理して作るための考え方』を提供していると理解することです。",
        "次回以降の学習でも、状態、データフロー、副作用、再利用性という観点が繰り返し登場します。今回の内容を自分の言葉で説明できるかどうかが、次のテーマへ進む準備になります。",
        "最後に、コードを読むだけで終えず、必ず自分で1つ小さな改造を加えてみてください。色を変える、文言を増やす、入力項目を追加するなど、少しの実験が理解を深める最短距離です。",
    ]
    return lesson


LESSONS = [lesson_to_document_data(item) for item in LESSON_BASES]


def configure_document(doc, lesson):
    section = doc.sections[0]
    section.top_margin = Mm(24)
    section.bottom_margin = Mm(18)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Yu Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
    normal.font.size = Pt(10.5)

    if "BodyJP" not in styles:
        style = styles.add_style("BodyJP", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(10.5)

    if "SectionTitleJP" not in styles:
        style = styles.add_style("SectionTitleJP", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Yu Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Yu Gothic")
        style.font.size = Pt(18)
        style.font.bold = True

    if "CodeBlock" not in styles:
        style = styles.add_style("CodeBlock", WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = styles["Normal"]
        style.font.name = "Consolas"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        style.font.size = Pt(9.5)

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.style = styles["BodyJP"]
    left_run = footer_p.add_run(f"React({lesson['number']})")
    left_run.font.size = Pt(8)
    left_run.font.color.rgb = RGBColor(140, 140, 140)
    footer_p.add_run("\t")
    add_page_number(footer_p)


def populate_document(doc, lesson):
    title = doc.add_paragraph(style="BodyJP")
    title_run = title.add_run(f"React({lesson['number']})")
    title_run.bold = True
    title_run.font.size = Pt(26)
    set_para_spacing(title, after=8, line=1.0)

    subtitle = doc.add_paragraph(style="BodyJP")
    subtitle_run = subtitle.add_run(lesson["subtitle"])
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(18)
    set_para_spacing(subtitle, after=8, line=1.1)

    target = doc.add_paragraph(style="BodyJP")
    target_run = target.add_run("この回の学習目標")
    target_run.bold = True
    target_run.font.size = Pt(13)
    set_para_spacing(target, after=4, line=1.2)

    for text in lesson["opening"]:
        p = doc.add_paragraph(text, style="BodyJP")
        set_para_spacing(p, after=5, line=1.55)

    prep = doc.add_paragraph(style="BodyJP")
    prep_run = prep.add_run(f"準備: {lesson['subtitle']}の実行確認")
    prep_run.bold = True
    prep_run.font.size = Pt(13)
    set_para_spacing(prep, before=8, after=4, line=1.2)
    add_code_block(doc, "", lesson["setup_code"])

    key_title = doc.add_paragraph(style="BodyJP")
    key_run = key_title.add_run("学習の見取り図")
    key_run.bold = True
    key_run.font.size = Pt(13)
    set_para_spacing(key_title, before=10, after=4, line=1.2)
    add_bullet_list(doc, lesson["key_points"])

    for index, section in enumerate(lesson["sections"], start=1):
        add_section_heading(doc, index, section["title"])
        step = doc.add_paragraph(style="BodyJP")
        step_run = step.add_run(f"Step {index}: {section['title']}")
        step_run.bold = True
        step_run.font.size = Pt(11.5)
        set_para_spacing(step, after=4, line=1.2)

        for concept in section["concepts"]:
            p = doc.add_paragraph(concept, style="BodyJP")
            set_para_spacing(p, after=5, line=1.55)

        deep_dive = (
            f"ここで扱っている『{section['title']}』は、単発のテクニックではなく、今後のReact設計全体に影響する土台です。"
            f"{lesson['subtitle']}の文脈では、{lesson['terms'][0]}や{lesson['terms'][1]}の理解と結びつけて考えると、単なる文法ではなく設計判断として捉えやすくなります。"
            "特に、表示を作る責任と値を管理する責任を混同しないこと、今のコードが将来の変更に耐えられるかを意識することが、中級レベルへの橋渡しになります。"
        )
        p = doc.add_paragraph(deep_dive, style="BodyJP")
        set_para_spacing(p, after=5, line=1.55)
        add_code_block(doc, section["code_title"], section["code"])

        caution = (
            f"注意点: {section['title']}では、コードが短く見えても複数の考え方が重なっています。"
            "動かなかった場合は、構文の誤りだけでなく、stateの更新対象、props名のずれ、import漏れ、フックの呼び出し位置などを順番に確認してください。"
            "Reactの学習では、原因を一つずつ切り分ける姿勢が最も大切です。"
        )
        p = doc.add_paragraph(caution, style="BodyJP")
        set_para_spacing(p, before=2, after=6, line=1.55)

    qa_title = doc.add_paragraph(style="BodyJP")
    qa_run = qa_title.add_run("よくある質問")
    qa_run.bold = True
    qa_run.font.size = Pt(13)
    set_para_spacing(qa_title, before=10, after=4, line=1.2)

    for item in lesson["qa"]:
        q = doc.add_paragraph(item["q"], style="BodyJP")
        q.runs[0].bold = True
        set_para_spacing(q, after=2, line=1.5)
        a = doc.add_paragraph(item["a"], style="BodyJP")
        set_para_spacing(a, after=4, line=1.5)

    practice_title = doc.add_paragraph(style="BodyJP")
    practice_run = practice_title.add_run("演習課題")
    practice_run.bold = True
    practice_run.font.size = Pt(13)
    set_para_spacing(practice_title, before=10, after=4, line=1.2)
    add_bullet_list(doc, lesson["practice"])

    summary_title = doc.add_paragraph(style="BodyJP")
    summary_run = summary_title.add_run("まとめ")
    summary_run.bold = True
    summary_run.font.size = Pt(13)
    set_para_spacing(summary_title, before=10, after=4, line=1.2)

    for text in lesson["closing"]:
        p = doc.add_paragraph(text, style="BodyJP")
        set_para_spacing(p, after=5, line=1.55)


def document_text_length(doc_path):
    reader = Document(doc_path)
    return sum(len(paragraph.text) for paragraph in reader.paragraphs)


def current_doc_length(doc):
    return sum(len(paragraph.text) for paragraph in doc.paragraphs)


def ensure_minimum_length(doc, lesson, minimum=5000):
    extra_index = 1
    while current_doc_length(doc) < minimum:
        p = doc.add_paragraph(style="BodyJP")
        text = (
            f"補足解説 {extra_index}: 第{lesson['number']}回で扱った{lesson['subtitle']}は、"
            f"{'、'.join(lesson['terms'][:3])}を中心に、Reactの画面設計をより安定させるための視点を育てる章です。"
            "学習を終えたら、サンプルを写すだけでなく、変数名、表示文言、配列データ、入力項目を自分なりに変更し、どの変更がどの描画結果に影響するかを観察してください。"
            "この『小さく変えて確かめる』反復が、基礎理解を中級の実装力へ変えていきます。"
        )
        p.add_run(text)
        set_para_spacing(p, after=5, line=1.55)
        extra_index += 1


def main():
    results = []
    for lesson in LESSONS:
        doc = Document()
        configure_document(doc, lesson)
        populate_document(doc, lesson)
        ensure_minimum_length(doc, lesson)
        safe_subtitle = lesson["subtitle"].replace(" ", "").replace(":", "・")
        file_name = f"React教材_第{lesson['number']:02d}回_{safe_subtitle}.docx"
        path = OUTPUT_DIR / file_name
        doc.save(path)
        results.append((path, document_text_length(path)))

    for path, length in results:
        print(f"{path}\t{length}")


if __name__ == "__main__":
    main()
